"""
writer_agent.py

RAG(기획서, 태스크)와 query를 바탕으로 문서 초안을 작성하고
.md / .docx / .pdf 형식으로 저장한다.

입력:
- query:      작성 요청 (어떤 문서를 쓸지)
- task_ids:   특정 태스크 ID 지정 (None이면 RAG 전체 검색)
- output_fmt: "md" | "docx" | "pdf" (기본 "md")
- filename:   저장 파일명 (확장자 제외, 기본 자동 생성)
- model_name: 사용할 LLM

출력:
- 작성된 문서 내용 (markdown 텍스트)
- 저장된 파일 경로
"""

import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from langchain_core.messages import HumanMessage

from app.core.utils import get_llm
from app.core.prompts import WRITER_PROMPT
from app.db.vector_store import get_rag_store, get_memory_store
from app.db.postgres import get_setting
from app.core.logger import get_logger

logger = get_logger("writer_agent")

_SUPPORTED_FORMATS = {"md", "docx", "pdf"}


async def run_writer(
    query: str,
    task_ids: list[int] | None = None,
    output_fmt: str = "md",
    filename: str | None = None,
    session_id: str = "default",
    model_name: str | None = None,
) -> dict:
    """
    문서를 작성하고 파일로 저장한다.

    Returns:
        {
            "content":   str,   # 마크다운 본문
            "filepath":  str,   # 저장된 파일 전체 경로
            "filename":  str,   # 파일명
            "format":    str,   # 실제 저장 포맷
        }
    """
    if output_fmt not in _SUPPORTED_FORMATS:
        raise ValueError(f"지원하지 않는 형식: {output_fmt}. 가능: {_SUPPORTED_FORMATS}")

    # 1. 컨텍스트 수집
    context = await _gather_context(query, task_ids)
    logger.info(f"컨텍스트 수집 완료: {len(context)}자")

    # 2. LLM으로 마크다운 문서 작성
    content = await _write(query, context, model_name)
    logger.info(f"문서 작성 완료: {len(content)}자")

    # 3. 파일 저장
    filepath = await _save(content, output_fmt, filename)
    saved_filename = Path(filepath).name
    logger.info(f"파일 저장 완료: {filepath}")

    return {
        "content":  content,
        "filepath": filepath,
        "filename": saved_filename,
        "format":   output_fmt,
    }


# ── 컨텍스트 수집 ─────────────────────────────────────────────

async def _gather_context(query: str, task_ids: list[int] | None) -> str:
    """RAG + 메모리에서 관련 컨텍스트를 수집한다."""
    parts = []

    try:
        rag_store = get_rag_store()

        if task_ids:
            # 특정 태스크 ID 지정 시 해당 태스크만 검색
            all_docs = rag_store.similarity_search(query, k=50)
            task_docs = [
                doc for doc in all_docs
                if doc.metadata.get("task_id") in task_ids
                or doc.metadata.get("source") in ("plan", "planner")
            ]
            docs = task_docs[:10] if task_docs else all_docs[:5]
        else:
            docs = rag_store.similarity_search(query, k=8)

        if docs:
            rag_text = "\n\n".join([
                f"[{doc.metadata.get('section_type') or doc.metadata.get('task_type') or 'doc'}] "
                f"{doc.page_content}"
                for doc in docs
            ])
            parts.append(f"=== 기획서 / 태스크 컨텍스트 ===\n{rag_text}")
    except Exception as e:
        logger.warning(f"RAG 컨텍스트 수집 실패: {e}")

    try:
        mem_store = get_memory_store()
        mem_docs = mem_store.similarity_search(query, k=3)
        if mem_docs:
            mem_text = "\n\n".join([doc.page_content for doc in mem_docs])
            parts.append(f"=== 과거 대화 컨텍스트 ===\n{mem_text}")
    except Exception as e:
        logger.warning(f"메모리 컨텍스트 수집 실패: {e}")

    return "\n\n".join(parts) if parts else "없음"


# ── LLM 문서 작성 ─────────────────────────────────────────────

async def _write(query: str, context: str, model_name: str | None) -> str:
    """LLM으로 마크다운 문서를 작성한다."""
    llm = get_llm(model_name)
    prompt = WRITER_PROMPT.format(query=query, context=context)

    response = await llm.ainvoke([HumanMessage(content=prompt)])
    content = response.content.strip()

    # 혹시 붙은 ```markdown 펜스 제거
    content = re.sub(r"^```(?:markdown)?\s*", "", content)
    content = re.sub(r"\s*```$", "", content)

    return content


# ── 파일 저장 ─────────────────────────────────────────────────

async def _save(content: str, fmt: str, filename: str | None) -> str:
    """마크다운 콘텐츠를 지정 형식으로 워크스페이스에 저장한다."""
    workspace = await _get_workspace()
    os.makedirs(workspace, exist_ok=True)

    stem = filename or _auto_filename(content)
    # 파일명에 허용되지 않는 문자 제거
    stem = re.sub(r'[\\/*?:"<>|]', "_", stem)

    if fmt == "md":
        return _save_md(content, workspace, stem)
    elif fmt == "docx":
        return _save_docx(content, workspace, stem)
    elif fmt == "pdf":
        return _save_pdf(content, workspace, stem)


def _save_md(content: str, workspace: str, stem: str) -> str:
    path = os.path.join(workspace, f"{stem}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def _save_docx(content: str, workspace: str, stem: str) -> str:
    """마크다운 → docx 변환. pandoc 우선, 없으면 python-docx 폴백."""
    path = os.path.join(workspace, f"{stem}.docx")

    # pandoc 시도
    if _pandoc_available():
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md",
                                         delete=False, encoding="utf-8") as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            subprocess.run(
                ["pandoc", tmp_path, "-o", path, "--from=markdown"],
                check=True, capture_output=True,
            )
            return path
        except subprocess.CalledProcessError as e:
            logger.warning(f"pandoc 변환 실패, python-docx 폴백: {e.stderr.decode()}")
        finally:
            os.unlink(tmp_path)

    # python-docx 폴백
    return _save_docx_python(content, path)


def _save_docx_python(content: str, path: str) -> str:
    """python-docx로 마크다운 → docx 변환 (헤딩/본문/bullet 기본 지원)."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx가 필요합니다. `pip install python-docx`")

    doc = DocxDocument()

    for line in content.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    doc.save(path)
    return path


def _save_pdf(content: str, workspace: str, stem: str) -> str:
    """마크다운 → pdf 변환. pandoc 우선, 없으면 reportlab 폴백."""
    path = os.path.join(workspace, f"{stem}.pdf")

    if _pandoc_available():
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md",
                                         delete=False, encoding="utf-8") as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            subprocess.run(
                ["pandoc", tmp_path, "-o", path,
                 "--from=markdown", "--pdf-engine=weasyprint"],
                check=True, capture_output=True,
            )
            return path
        except subprocess.CalledProcessError:
            pass
        finally:
            os.unlink(tmp_path)

    # reportlab 폴백
    return _save_pdf_reportlab(content, path)


def _save_pdf_reportlab(content: str, path: str) -> str:
    """reportlab으로 마크다운 텍스트를 PDF로 저장한다."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 한글 폰트 등록 (없으면 기본 폰트 사용)
    font_regular = "Helvetica"
    font_bold    = "Helvetica-Bold"
    nanum_path   = "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"
    nanum_bold   = "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf"
    if os.path.exists(nanum_path):
        pdfmetrics.registerFont(TTFont("NanumGothic", nanum_path))
        pdfmetrics.registerFont(TTFont("NanumGothicBold", nanum_bold))
        font_regular = "NanumGothic"
        font_bold    = "NanumGothicBold"

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", fontName=font_bold,   fontSize=16, spaceAfter=8,  textColor=colors.HexColor("#1a1a2e"))
    h2 = ParagraphStyle("H2", fontName=font_bold,   fontSize=13, spaceAfter=6,  spaceBefore=12, textColor=colors.HexColor("#16213e"))
    h3 = ParagraphStyle("H3", fontName=font_bold,   fontSize=11, spaceAfter=4,  spaceBefore=8)
    body = ParagraphStyle("Body", fontName=font_regular, fontSize=10, leading=16, spaceAfter=4)
    bullet = ParagraphStyle("Bullet", fontName=font_regular, fontSize=10, leading=16,
                             leftIndent=14, spaceAfter=3)

    story = []
    for line in content.splitlines():
        # HTML 특수문자 이스케이프
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if safe.startswith("### "):
            story.append(Paragraph(safe[4:], h3))
        elif safe.startswith("## "):
            story.append(Paragraph(safe[3:], h2))
        elif safe.startswith("# "):
            story.append(Paragraph(safe[2:], h1))
        elif safe.startswith("- ") or safe.startswith("* "):
            story.append(Paragraph(f"• {safe[2:]}", bullet))
        elif safe.strip() == "":
            story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Paragraph(safe, body))

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )
    doc.build(story)
    return path


# ── 헬퍼 ──────────────────────────────────────────────────────

async def _get_workspace() -> str:
    """설정된 워크스페이스 경로를 반환. 없으면 /tmp/research_workspace."""
    try:
        path = await get_setting("workspace_path")
        if path:
            return path
    except Exception:
        pass
    return "/tmp/research_workspace"


def _auto_filename(content: str) -> str:
    """문서 첫 줄(# 제목)에서 파일명을 생성한다."""
    for line in content.splitlines():
        if line.startswith("# "):
            name = line[2:].strip()[:40]
            name = re.sub(r"\s+", "_", name)
            return name
    return f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"


def _pandoc_available() -> bool:
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False