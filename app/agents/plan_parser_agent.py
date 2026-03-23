"""
plan_parser_agent.py
"""

import json
import re
import io
from datetime import datetime

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.utils import get_llm
from app.core.prompts import CLASSIFY_PROMPT, PARSE_PROMPTS
from app.db.vector_store import get_rag_store, deactivate_documents_by_filename
from app.core.logger import get_logger
from sqlalchemy import text

logger = get_logger("plan_parser_agent")

_HEADER_LENGTH = 500  # 문서 타입 분류에 사용할 앞부분 길이

# 문서 타입별 청킹 전략
_CHUNK_STRATEGIES = {
    "PLAN": {
        "chunk_size": 600,
        "chunk_overlap": 80,
        "separators": ["\n\n", "\n", "。", ". ", " ", ""],
    },
    "REPORT": {
        "chunk_size": 800,
        "chunk_overlap": 100,
        "separators": ["\n\n", "\n", "。", ". ", " ", ""],
    },
    "ARTICLE": {
        "chunk_size": 500,
        "chunk_overlap": 50,
        "separators": ["\n\n", "\n", ". ", " ", ""],
    },
    "LEGAL": {
        "chunk_size": 400,
        "chunk_overlap": 60,
        "separators": ["\n\n", "\n", ". ", " ", ""],
    },
    "RESUME": {
        "chunk_size": 400,
        "chunk_overlap": 40,
        "separators": ["\n\n", "\n", ". ", " ", ""],
    },
    "GENERAL": {
        "chunk_size": 500,
        "chunk_overlap": 50,
        "separators": ["\n\n", "\n", ". ", " ", ""],
    },
}

async def parse_and_ingest_plan(
    filename: str,
    file_bytes: bytes,
    content_type: str,
    model_name: str | None = None,
) -> dict:
    # 1. 텍스트 추출 (기존 코드)
    text_content = _extract_text(filename, file_bytes)
    
    # 2. 문서 타입 분류 및 섹션 파싱 (기존 코드)
    doc_type = await _classify_document(text_content[:_HEADER_LENGTH], model_name)
    parsed = await _parse_sections(text_content[:4000], doc_type, model_name)
    
    # [추가] 3. 신규 저장 전 기존 데이터 비활성화
    await deactivate_documents_by_filename(filename)

    # 4. Document 생성 시 is_active 메타데이터 주입
    title = parsed.get("title", filename)
    sections = parsed.get("sections", [])
    
    # _build_documents 함수를 수정하거나, 생성된 docs의 metadata를 직접 수정
    docs = _build_documents(filename, title, doc_type, sections)
    
    for doc in docs:
        doc.metadata["is_active"] = True  # 새 데이터는 활성 상태
        doc.metadata["version"] = datetime.utcnow().isoformat() # 버전 기록

    # 5. RAG 저장
    store = get_rag_store()
    store.add_documents(docs)

    return { "title": title, "doc_type": doc_type, "sections": len(sections), "chunks": len(docs) }

# ── 문서 타입 분류 ────────────────────────────────────────────

async def _classify_document(header: str, model_name: str | None) -> str:
    """문서 앞부분만 읽어 타입을 분류한다. 실패 시 GENERAL 반환."""
    llm = get_llm(model_name)
    prompt = CLASSIFY_PROMPT.format(header=header)

    try:
        response = await llm.ainvoke([HumanMessage(content=prompt)])
        doc_type = response.content.strip().upper()

        # 유효한 타입인지 확인, 아니면 GENERAL로 폴백
        valid_types = {"PLAN", "REPORT", "ARTICLE", "LEGAL", "RESUME", "GENERAL"}
        if doc_type not in valid_types:
            logger.warning(f"알 수 없는 문서 타입: {doc_type} → GENERAL로 폴백")
            return "GENERAL"

        return doc_type

    except Exception as e:
        logger.error(f"문서 타입 분류 실패: {e} → GENERAL로 폴백")
        return "GENERAL"


# ── LLM 섹션 파싱 ─────────────────────────────────────────────

async def _parse_sections(text: str, doc_type: str, model_name: str | None) -> dict:
    """타입별 프롬프트로 섹션을 파싱한다. 실패 시 빈 dict 반환."""
    llm = get_llm(model_name)

    # 타입별 프롬프트 선택 — 없으면 GENERAL로 폴백
    prompt_template = PARSE_PROMPTS.get(doc_type, PARSE_PROMPTS["GENERAL"])
    prompt = prompt_template.format(content=text)

    try:
        response = await llm.ainvoke([HumanMessage(content=prompt)])
        raw = response.content.strip()

        # ```json ... ``` 펜스 제거
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        parsed = json.loads(raw)

        # title 빈 문자열 폴백
        if not parsed.get("title"):
            parsed["title"] = text.strip().splitlines()[0].strip()

        return parsed

    except json.JSONDecodeError as e:
        logger.error(f"섹션 파싱 JSON 오류: {e} / raw: {raw[:200]}")
        return {}
    except Exception as e:
        logger.error(f"섹션 파싱 실패: {e}")
        return {}


# ── Document 빌드 ─────────────────────────────────────────────

def _get_splitter(doc_type: str) -> RecursiveCharacterTextSplitter:
    """문서 타입별 최적화된 텍스트 스플리터를 반환한다."""
    strategy = _CHUNK_STRATEGIES.get(doc_type, _CHUNK_STRATEGIES["GENERAL"])
    return RecursiveCharacterTextSplitter(
        chunk_size=strategy["chunk_size"],
        chunk_overlap=strategy["chunk_overlap"],
        separators=strategy["separators"],
        length_function=len,
    )

def _build_documents(
    filename: str,
    title: str,
    doc_type: str,
    sections: list[dict],
) -> list[Document]:
    """
    섹션 리스트를 langchain Document로 변환한다.
    문서 타입별 청킹 전략을 적용한다.
    """
    splitter = _get_splitter(doc_type)
    docs = []

    for i, section in enumerate(sections):
        section_type    = section.get("type", "other")
        section_heading = section.get("heading", f"섹션 {i + 1}")
        content         = section.get("content", "")

        if isinstance(content, (dict, list)):
            import json
            logger.warning(f"섹션 {i+1} content가 {type(content).__name__} 타입 — 문자열 변환")
            content = json.dumps(content, ensure_ascii=False)

        content = str(content).strip()
        if not content:
            continue

        # 섹션 제목을 content 앞에 붙여 검색 정확도 향상
        enriched_content = f"[{section_heading}]\n{content}"

        base_metadata = {
            "filename":        filename,
            "plan_title":      title,
            "doc_type":        doc_type,
            "section_type":    section_type,
            "section_heading": section_heading,
            "source":          "plan",
        }

        # 섹션이 짧으면 단일 Document
        if len(enriched_content) <= splitter._chunk_size:
            docs.append(Document(
                page_content=enriched_content,
                metadata=base_metadata,
            ))
            continue

        # 길면 타입별 스플리터로 분할
        chunks = splitter.split_text(enriched_content)
        for j, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={**base_metadata, "chunk_index": j},
            ))

    return docs

# ── 텍스트 추출 ───────────────────────────────────────────────

def _extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    if lower.endswith((".md", ".txt")):
        return file_bytes.decode("utf-8", errors="ignore")

    raise ValueError(f"지원하지 않는 파일 형식입니다: {filename}")


def _extract_pdf(file_bytes: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx가 필요합니다. `pip install python-docx`")

    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())